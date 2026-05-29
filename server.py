#!/usr/bin/env python3
"""
アビリティスタッフ HP サーバー
- HTMLファイル・静的ファイル配信
- Anthropic APIプロキシ（AIマッチング）
- PDF/Word/Excel テキスト抽出
- PostgreSQL DB接続（jobmatchと同じDB）
"""

import http.server
import json
import urllib.request
import urllib.error
import os
import io

PORT = int(os.environ.get('PORT', 8000))
API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')

# DB接続URL（RenderではDATABASE_URL、ローカルではEXTERNAL_DATABASE_URLを使用）
DATABASE_URL = os.environ.get('DATABASE_URL') or os.environ.get('EXTERNAL_DATABASE_URL', '')

# ===== DB接続 =====

def get_db_connection():
    try:
        import psycopg2
        return psycopg2.connect(DATABASE_URL)
    except Exception as e:
        print(f'DB接続エラー: {e}')
        return None

def load_jobs_from_db():
    """DBから求人データを軽量版で取得（jobmatchと同じ構造）"""
    conn = get_db_connection()
    if not conn:
        return None
    try:
        cur = conn.cursor()
        cur.execute("""
            SELECT j.id, j.company, j.category, j.subcategory, j.title,
                   LEFT(j.description, 120), j.salary, j.employment,
                   j.prefecture, j.tags, j.updated,
                   a.age, a.job_change_count, a.gender, a.note,
                   LEFT(j.requirements, 200)
            FROM jobs j
            LEFT JOIN job_agent_info a ON j.id = a.job_id
            ORDER BY j.id;
        """)
        rows = cur.fetchall()
        jobs = []
        for row in rows:
            job = {
                'id': str(row[0]),
                'company': row[1] or '',
                'category': row[2] or '',
                'subcategory': row[3] or '',
                'title': row[4] or '',
                'description': row[5] or '',
                'salary': row[6] or '',
                'employment': row[7] or '',
                'prefecture': row[8] or '',
                'tags': json.loads(row[9]) if row[9] else [],
                'updated': str(row[10]) if row[10] else '',
                'agent_condition': {
                    'age': row[11] or '',
                    'job_change_count': row[12] or '',
                    'gender': row[13] or '',
                    'note': row[14] or '',
                },
                'requirements': row[15] or '',
            }
            jobs.append(job)
        cur.close()
        conn.close()
        return jobs
    except Exception as e:
        print(f'DB読み込みエラー: {e}')
        if conn:
            conn.close()
        return None

def load_job_detail_from_db(job_id):
    """DBから1件の詳細データを取得（jobmatchと同じ構造）"""
    conn = get_db_connection()
    if not conn:
        return None
    try:
        cur = conn.cursor()
        cur.execute("""
            SELECT j.id, j.company, j.category, j.subcategory, j.title,
                   j.description, j.requirements, j.salary, j.employment,
                   j.prefecture, j.tags, j.updated,
                   a.gender, a.age, a.foreign_national, a.job_change_count,
                   a.headcount, a.fee, a.fee_definition, a.refund, a.note, a.other
            FROM jobs j
            LEFT JOIN job_agent_info a ON j.id = a.job_id
            WHERE j.id = %s;
        """, (job_id,))
        row = cur.fetchone()
        cur.close()
        conn.close()
        if not row:
            return None
        job = {
            'id': str(row[0]),
            'company': row[1] or '',
            'category': row[2] or '',
            'subcategory': row[3] or '',
            'title': row[4] or '',
            'description': row[5] or '',
            'requirements': row[6] or '',
            'salary': row[7] or '',
            'employment': row[8] or '',
            'prefecture': row[9] or '',
            'tags': json.loads(row[10]) if row[10] else [],
            'updated': str(row[11]) if row[11] else '',
            'agent': {
                'gender': row[12] or '',
                'age': row[13] or '',
                'foreign_national': row[14] or '',
                'job_change_count': row[15] or '',
                'headcount': row[16] or '',
                'fee': row[17] or '',
                'fee_definition': row[18] or '',
                'refund': row[19] or '',
                'note': row[20] or '',
                'other': row[21] or '',
            }
        }
        return job
    except Exception as e:
        print(f'DB詳細読み込みエラー: {e}')
        if conn:
            conn.close()
        return None

# ===== テキスト抽出 =====

def extract_pdf_text(file_bytes):
    try:
        import fitz
        import re
        doc = fitz.open(stream=file_bytes, filetype='pdf')
        texts = []
        for page in doc:
            blocks = page.get_text('blocks')
            for block in blocks:
                if block[6] == 0:
                    text = block[4].strip()
                    if text:
                        texts.append(text)
        doc.close()
        full_text = '\n'.join(texts)
        full_text = re.sub(r'([^\n])【', r'\1\n【', full_text)
        full_text = re.sub(r'([^\n\d])(\d{4}年)', r'\1\n\2', full_text)
        return full_text
    except Exception as e:
        return f'[PDF抽出エラー: {str(e)}]'

def extract_docx_text(file_bytes):
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        texts = []
        seen = set()
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and text not in seen:
                seen.add(text)
                texts.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in seen:
                        seen.add(text)
                        texts.append(text)
        return '\n'.join(texts)
    except Exception:
        pass
    try:
        import docx2txt
        import tempfile
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as f:
            f.write(file_bytes)
            tmp_path = f.name
        text = docx2txt.process(tmp_path)
        os.unlink(tmp_path)
        return text or ''
    except Exception as e:
        return f'[Word抽出エラー: {str(e)}]'

def extract_xlsx_text(file_bytes):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        texts = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and str(cell.value).strip():
                        texts.append(str(cell.value).strip())
        wb.close()
        return '\n'.join(texts)
    except Exception:
        pass
    try:
        import xlrd
        wb = xlrd.open_workbook(file_contents=file_bytes)
        texts = []
        for sheet in wb.sheets():
            for row in range(sheet.nrows):
                for col in range(sheet.ncols):
                    val = str(sheet.cell_value(row, col)).strip()
                    if val and val != '0.0':
                        texts.append(val)
        return '\n'.join(texts)
    except Exception as e:
        return f'[Excel抽出エラー: {str(e)}]'

# ===== マルチパートパーサー =====

def parse_multipart(body, boundary):
    b = boundary.encode()
    delimiter = b'\r\n--' + b
    if body.startswith(b'--' + b):
        body = body[len(b'--' + b):]
    parts = body.split(delimiter)

    file_bytes = None
    filename = ''
    text_data = ''

    for part in parts:
        if b'Content-Disposition' not in part:
            continue
        if b'\r\n\r\n' not in part:
            continue
        header_raw, content = part.split(b'\r\n\r\n', 1)
        header_str = header_raw.decode('utf-8', errors='ignore')

        while content.endswith(b'\r\n'):
            content = content[:-2]
        while content.endswith(b'--'):
            content = content[:-2]
        while content.endswith(b'\r\n'):
            content = content[:-2]

        if 'name="file"' in header_str:
            file_bytes = content
            for seg in header_str.replace('\r\n', ';').split(';'):
                seg = seg.strip()
                if seg.startswith('filename='):
                    filename = seg[9:].strip().strip('"')

        elif 'name="text"' in header_str:
            text_data = content.decode('utf-8', errors='ignore')

    return file_bytes, filename, text_data

# ===== HTTPハンドラー =====

class AbilityStaffHandler(http.server.SimpleHTTPRequestHandler):

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'GET, POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type, x-api-key, anthropic-version')
        self.end_headers()

    def do_GET(self):
        # ヘルスチェック
        if self.path == '/api/health':
            status = {'server': 'OK', 'db': 'NG'}
            conn = get_db_connection()
            if conn:
                status['db'] = 'OK'
                conn.close()
            self.send_json(status)
            return

        # 求人一覧API
        if self.path == '/api/jobs':
            if not DATABASE_URL:
                self.send_error(503, 'DATABASE_URL not set')
                return
            jobs = load_jobs_from_db()
            if jobs is None:
                self.send_error(500, 'DB connection failed')
                return
            self.send_json(jobs)
            return

        # 求人詳細API
        if self.path.startswith('/api/jobs/detail/'):
            job_id = self.path.replace('/api/jobs/detail/', '').split('?')[0]
            if not DATABASE_URL:
                self.send_error(503, 'DATABASE_URL not set')
                return
            job = load_job_detail_from_db(job_id)
            if job is None:
                self.send_error(404, 'Job not found')
                return
            self.send_json(job)
            return

        # その他は静的ファイルとして配信
        return super().do_GET()

    def do_POST(self):
        if self.path == '/api/claude':
            self.proxy_to_claude()
        elif self.path == '/api/extract':
            self.extract_text()
        else:
            self.send_error(404)

    def send_json(self, data):
        response = json.dumps(data, ensure_ascii=False).encode('utf-8')
        self.send_response(200)
        self.send_header('Content-Type', 'application/json; charset=utf-8')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        self.wfile.write(response)

    def proxy_to_claude(self):
        """Anthropic APIへのプロキシ"""
        try:
            length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(length)

            req = urllib.request.Request(
                'https://api.anthropic.com/v1/messages',
                data=body,
                method='POST'
            )
            req.add_header('Content-Type', 'application/json')
            req.add_header('x-api-key', API_KEY)
            req.add_header('anthropic-version', '2023-06-01')

            with urllib.request.urlopen(req) as resp:
                response_body = resp.read()
                self.send_response(200)
                self.send_header('Content-Type', 'application/json')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(response_body)
        except urllib.error.HTTPError as e:
            error_body = e.read()
            self.send_response(e.code)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(error_body)
        except Exception as e:
            self.send_error(500, str(e))

    def extract_text(self):
        """ファイルからテキストを抽出して返す"""
        try:
            content_type = self.headers.get('Content-Type', '')
            length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(length)

            boundary = None
            for item in content_type.split(';'):
                item = item.strip()
                if item.lower().startswith('boundary='):
                    boundary = item[9:].strip().strip('"')
                    break

            if not boundary:
                self.send_error(400, 'No boundary')
                return

            file_bytes, filename, text_data = parse_multipart(body, boundary)

            # テキスト直接入力の場合
            if text_data and not file_bytes:
                self.send_json({'text': text_data, 'filename': 'direct_input'})
                return

            # ファイルアップロードの場合
            if file_bytes:
                fn = filename.lower()
                if fn.endswith('.pdf'):
                    text = extract_pdf_text(file_bytes)
                elif fn.endswith('.docx') or fn.endswith('.doc'):
                    text = extract_docx_text(file_bytes)
                elif fn.endswith('.xlsx') or fn.endswith('.xls'):
                    text = extract_xlsx_text(file_bytes)
                elif fn.endswith('.txt'):
                    text = file_bytes.decode('utf-8', errors='ignore')
                else:
                    text = f'[未対応のファイル形式: {filename}]'

                self.send_json({'text': text, 'filename': filename})
                return

            self.send_error(400, 'No file or text provided')

        except Exception as e:
            self.send_error(500, str(e))

    def log_message(self, format, *args):
        print(f'[{self.address_string()}] {format % args}')


if __name__ == '__main__':
    import socketserver
    print(f'アビリティスタッフ HP サーバー起動中: http://localhost:{PORT}')
    print(f'DB: {"接続設定あり" if DATABASE_URL else "未設定（DBなしで起動）"}')
    print('終了するには Ctrl+C を押してください')

    with socketserver.TCPServer(('', PORT), AbilityStaffHandler) as httpd:
        httpd.serve_forever()
